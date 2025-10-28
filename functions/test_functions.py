"""
Quick test script for Cloud Functions
Run this to verify basic functionality before deploying
"""

import sys
from pathlib import Path

# Add src to path
sys.path.insert(0, str(Path(__file__).parent / 'src'))

def test_imports():
    """Test that all modules can be imported"""
    print("Testing imports...")

    try:
        from src.documents.placeholders import replace_placeholders, find_placeholders
        print("  ✓ placeholders module")

        from src.documents.generate import generate_documents, generate_single_document
        print("  ✓ generate module")

        from src.documents.regenerate import regenerate_document
        print("  ✓ regenerate module")

        from src.templates.analyze import analyze_template
        print("  ✓ analyze module")

        from src.projects.create import create_project
        print("  ✓ create project module")

        from src.projects.update_status import update_project_status, VALID_STATUSES
        print("  ✓ update status module")

        from src.projects.variables import prepare_standard_variables
        print("  ✓ variables module")

        from src.utils.document_number import generate_document_number
        print("  ✓ document number module")

        print("\n✓ All imports successful!")
        return True

    except Exception as e:
        print(f"\n✗ Import failed: {e}")
        import traceback
        traceback.print_exc()
        return False


def test_document_number_generation():
    """Test document number generation"""
    print("\nTesting document number generation...")

    try:
        from src.utils.document_number import generate_document_number
        from datetime import datetime

        # Test with sample data (Oct 28, 2025)
        date = datetime(2025, 10, 28)
        counter = 1
        result = generate_document_number(date, counter)

        print(f"  Generated number: {result}")

        # Verify format (should be HIYESYYMDDNNN)
        # Example: HIYES25JBB001 (13 chars total)
        # 25 = Year 2025
        # J = October (10th month)
        # BB = Day 28
        # 001 = First document
        assert result.startswith('HIYES'), f"Expected prefix 'HIYES', got '{result[:5]}'"
        assert len(result) == 13, f"Expected length 13, got {len(result)}"

        # Test multiple counters
        test_cases = [
            (datetime(2025, 1, 1), 1, "HIYES25AAA001"),
            (datetime(2025, 10, 27), 1, "HIYES25JBA001"),
            (datetime(2025, 12, 31), 10, "HIYES25LBE010"),
        ]

        for test_date, test_counter, expected in test_cases:
            result = generate_document_number(test_date, test_counter)
            assert result == expected, f"Expected {expected}, got {result}"

        print("  ✓ Document number format correct")
        print("  ✓ All test cases passed")
        return True

    except Exception as e:
        print(f"  ✗ Test failed: {e}")
        import traceback
        traceback.print_exc()
        return False


def test_status_validation():
    """Test status validation"""
    print("\nTesting status validation...")

    try:
        from src.projects.update_status import VALID_STATUSES

        expected_statuses = [
            'draft',
            'in_progress',
            'paused',
            'pending_invoice',
            'pending_payment',
            'completed'
        ]

        print(f"  Valid statuses: {', '.join(VALID_STATUSES)}")

        for status in expected_statuses:
            assert status in VALID_STATUSES, f"Missing status: {status}"

        print("  ✓ All expected statuses present")
        return True

    except Exception as e:
        print(f"  ✗ Test failed: {e}")
        return False


def test_placeholder_regex():
    """Test placeholder detection"""
    print("\nTesting placeholder detection...")

    try:
        from src.documents.placeholders import find_placeholders
        from docx import Document

        # Create a test document in memory
        doc = Document()
        doc.add_paragraph("Project: {{project_name}}")
        doc.add_paragraph("Company: {{company_name}}")
        doc.add_paragraph("Price: NT$ {{price}}")

        placeholders = find_placeholders(doc)

        print(f"  Found placeholders: {', '.join(sorted(placeholders))}")

        expected = {'project_name', 'company_name', 'price'}
        assert placeholders == expected, f"Expected {expected}, got {placeholders}"

        print("  ✓ Placeholder detection working")
        return True

    except Exception as e:
        print(f"  ✗ Test failed: {e}")
        import traceback
        traceback.print_exc()
        return False


def main():
    """Run all tests"""
    print("=" * 60)
    print("Cloud Functions Unit Tests")
    print("=" * 60)

    tests = [
        test_imports,
        test_document_number_generation,
        test_status_validation,
        test_placeholder_regex,
    ]

    results = []
    for test in tests:
        results.append(test())

    print("\n" + "=" * 60)
    print(f"Results: {sum(results)}/{len(results)} tests passed")
    print("=" * 60)

    if all(results):
        print("\n✓ All tests passed! Functions are ready for deployment.")
        return 0
    else:
        print("\n✗ Some tests failed. Please fix the issues before deploying.")
        return 1


if __name__ == "__main__":
    sys.exit(main())
