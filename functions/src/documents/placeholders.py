"""
Placeholder Replacement Logic for Word Documents

Handles replacing {{variable}} placeholders in Word documents while
preserving formatting.
"""

from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import Table
from typing import Dict


def replace_placeholders(doc: Document, data: Dict[str, str]) -> None:
    """
    Replace all {{variable}} placeholders in the document with actual values.

    Searches in:
    - Paragraphs
    - Tables
    - Headers
    - Footers

    Args:
        doc: python-docx Document object
        data: Dictionary mapping variable names to values
    """

    # Replace in paragraphs
    for paragraph in doc.paragraphs:
        replace_in_paragraph(paragraph, data)

    # Replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_in_paragraph(paragraph, data)

    # Replace in headers and footers
    for section in doc.sections:
        # Header
        for paragraph in section.header.paragraphs:
            replace_in_paragraph(paragraph, data)

        # Footer
        for paragraph in section.footer.paragraphs:
            replace_in_paragraph(paragraph, data)


def replace_in_paragraph(paragraph: Paragraph, data: Dict[str, str]) -> None:
    """
    Replace placeholders in a single paragraph while preserving formatting.

    Handles cases where placeholders are split across multiple runs.

    Args:
        paragraph: Paragraph object
        data: Dictionary mapping variable names to values
    """

    # Get full text
    full_text = paragraph.text

    # Check if any placeholder exists
    has_placeholder = any(f"{{{{{key}}}}}" in full_text for key in data.keys())

    if not has_placeholder:
        return

    # Replace all placeholders
    new_text = full_text
    for key, value in data.items():
        placeholder = f"{{{{{key}}}}}"
        if placeholder in new_text:
            new_text = new_text.replace(placeholder, str(value))

    # If text changed, update the paragraph
    if new_text != full_text:
        # Simple approach: Replace text in first run, clear others
        # This preserves the first run's formatting
        if paragraph.runs:
            paragraph.runs[0].text = new_text

            # Clear remaining runs
            for i in range(1, len(paragraph.runs)):
                paragraph.runs[i].text = ''
        else:
            # No runs exist, add text directly
            paragraph.text = new_text


def find_placeholders(doc: Document) -> set:
    """
    Find all {{variable}} placeholders in the document.

    Args:
        doc: python-docx Document object

    Returns:
        Set of variable names found in the document
    """
    import re

    variables = set()
    pattern = r'\{\{([a-zA-Z0-9_]+)\}\}'

    # Search in paragraphs
    for paragraph in doc.paragraphs:
        matches = re.findall(pattern, paragraph.text)
        variables.update(matches)

    # Search in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    matches = re.findall(pattern, paragraph.text)
                    variables.update(matches)

    # Search in headers/footers
    for section in doc.sections:
        for paragraph in section.header.paragraphs:
            matches = re.findall(pattern, paragraph.text)
            variables.update(matches)

        for paragraph in section.footer.paragraphs:
            matches = re.findall(pattern, paragraph.text)
            variables.update(matches)

    return variables
