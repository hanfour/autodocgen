"""
Document Processor

Handle document template processing:
- Extract variables from documents
- Replace variables with values
- Generate output in multiple formats
"""

import re
import io
from typing import Dict, List, Optional, Tuple
from docx import Document
from docx.shared import Pt
from PyPDF2 import PdfReader
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch


def extract_variables(text: str) -> List[str]:
    """
    Extract variables from text using {{variable}} syntax

    Args:
        text: Text content to scan for variables

    Returns:
        List of unique variable names
    """
    pattern = r'\{\{([^}]+)\}\}'
    matches = re.findall(pattern, text)
    # Clean and deduplicate
    variables = [var.strip() for var in matches]
    return list(dict.fromkeys(variables))  # Preserve order, remove duplicates


def extract_variables_from_docx(file_path: str) -> Tuple[str, List[str]]:
    """
    Extract text and variables from Word document

    Args:
        file_path: Path to .docx file

    Returns:
        Tuple of (text_content, variables_list)
    """
    doc = Document(file_path)

    # Extract all text
    full_text = []
    for paragraph in doc.paragraphs:
        full_text.append(paragraph.text)

    text_content = '\n'.join(full_text)
    variables = extract_variables(text_content)

    return text_content, variables


def extract_variables_from_pdf(file_path: str) -> Tuple[str, List[str]]:
    """
    Extract text and variables from PDF

    Args:
        file_path: Path to .pdf file

    Returns:
        Tuple of (text_content, variables_list)
    """
    reader = PdfReader(file_path)

    # Extract all text
    full_text = []
    for page in reader.pages:
        full_text.append(page.extract_text())

    text_content = '\n'.join(full_text)
    variables = extract_variables(text_content)

    return text_content, variables


def replace_variables(text: str, values: Dict[str, str]) -> str:
    """
    Replace variables in text with provided values

    Args:
        text: Text with {{variable}} placeholders
        values: Dictionary mapping variable names to values

    Returns:
        Text with variables replaced
    """
    result = text
    for var_name, var_value in values.items():
        pattern = r'\{\{' + re.escape(var_name) + r'\}\}'
        result = re.sub(pattern, str(var_value), result)

    return result


def generate_docx(template_path: str, values: Dict[str, str], output_path: str):
    """
    Generate Word document from template with variable replacement

    Args:
        template_path: Path to template .docx file
        values: Variable values to replace
        output_path: Path for output file
    """
    # Load template
    doc = Document(template_path)

    # Replace variables in paragraphs
    for paragraph in doc.paragraphs:
        if '{{' in paragraph.text:
            # Get the original text
            original_text = paragraph.text
            # Replace variables
            new_text = replace_variables(original_text, values)
            # Update paragraph
            paragraph.text = new_text

    # Replace variables in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if '{{' in paragraph.text:
                        original_text = paragraph.text
                        new_text = replace_variables(original_text, values)
                        paragraph.text = new_text

    # Save output
    doc.save(output_path)


def generate_docx_from_text(text: str, values: Dict[str, str], output_path: str):
    """
    Generate Word document from plain text template

    Args:
        text: Template text with {{variables}}
        values: Variable values to replace
        output_path: Path for output file
    """
    # Create new document
    doc = Document()

    # Replace variables
    processed_text = replace_variables(text, values)

    # Add paragraphs
    for line in processed_text.split('\n'):
        doc.add_paragraph(line)

    # Save
    doc.save(output_path)


def generate_pdf_from_text(text: str, values: Dict[str, str], output_path: str):
    """
    Generate PDF from plain text template

    Args:
        text: Template text with {{variables}}
        values: Variable values to replace
        output_path: Path for output file
    """
    # Replace variables
    processed_text = replace_variables(text, values)

    # Create PDF
    doc = SimpleDocTemplate(output_path, pagesize=letter)

    # Styles
    styles = getSampleStyleSheet()
    story = []

    # Add content
    for line in processed_text.split('\n'):
        if line.strip():
            para = Paragraph(line, styles['Normal'])
            story.append(para)
            story.append(Spacer(1, 0.2 * inch))

    # Build PDF
    doc.build(story)


def generate_html(text: str, values: Dict[str, str]) -> str:
    """
    Generate HTML from text template

    Args:
        text: Template text with {{variables}}
        values: Variable values to replace

    Returns:
        HTML string
    """
    # Replace variables
    processed_text = replace_variables(text, values)

    # Convert to HTML (simple version)
    html_lines = []
    html_lines.append('<!DOCTYPE html>')
    html_lines.append('<html>')
    html_lines.append('<head>')
    html_lines.append('  <meta charset="UTF-8">')
    html_lines.append('  <style>')
    html_lines.append('    body { font-family: Arial, sans-serif; max-width: 800px; margin: 40px auto; padding: 20px; line-height: 1.6; }')
    html_lines.append('    p { margin-bottom: 1em; }')
    html_lines.append('  </style>')
    html_lines.append('</head>')
    html_lines.append('<body>')

    for line in processed_text.split('\n'):
        if line.strip():
            html_lines.append(f'  <p>{line}</p>')
        else:
            html_lines.append('  <br>')

    html_lines.append('</body>')
    html_lines.append('</html>')

    return '\n'.join(html_lines)
