from flask import Flask, render_template, request, redirect, url_for, flash, send_from_directory
import os
import json
from docx import Document
from werkzeug.utils import secure_filename

app = Flask(__name__)
app.secret_key = 'your_secret_key'

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
UPLOAD_FOLDER = os.path.join(BASE_DIR, 'uploads')
CONFIG_FOLDER = os.path.join(BASE_DIR, 'config')

os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(CONFIG_FOLDER, exist_ok=True)

ALLOWED_EXTENSIONS = {'docx'}

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

@app.route('/')
def index():
    # 列出所有已上傳模板
    templates_list = []
    for filename in os.listdir(UPLOAD_FOLDER):
        if filename.endswith('.docx'):
            template_name = os.path.splitext(filename)[0]
            templates_list.append({
                "filename": filename,
                "template_name": template_name
            })
    return render_template('index.html', templates_list=templates_list)

@app.route('/upload_template', methods=['GET', 'POST'])
def upload_template():
    if request.method == 'POST':
        if 'file' not in request.files:
            flash('沒有檔案部分')
            return redirect(request.url)
        file = request.files['file']
        if file.filename == '':
            flash('未選擇檔案')
            return redirect(request.url)
        if file and allowed_file(file.filename):
            filename = secure_filename(file.filename)
            save_path = os.path.join(UPLOAD_FOLDER, filename)
            file.save(save_path)

            # 解析佔位符
            placeholders = extract_placeholders(save_path)
            # 儲存佔位符設定
            config_path = os.path.join(CONFIG_FOLDER, f"{os.path.splitext(filename)[0]}_placeholders.json")
            with open(config_path, 'w', encoding='utf-8') as f:
                json.dump({"placeholders": list(placeholders)}, f, ensure_ascii=False, indent=2)

            flash('模板上傳成功，已解析佔位符')
            return redirect(url_for('index'))
        else:
            flash('只允許上傳 docx 檔案')
            return redirect(request.url)
    return render_template('upload_template.html')

@app.route('/uploads/<filename>')
def uploaded_file(filename):
    return send_from_directory(UPLOAD_FOLDER, filename)

@app.route('/delete_template/<template_name>', methods=['POST'])
def delete_template(template_name):
    filename = f"{template_name}.docx"
    file_path = os.path.join(UPLOAD_FOLDER, filename)
    placeholders_path = os.path.join(CONFIG_FOLDER, f"{template_name}_placeholders.json")
    fields_path = os.path.join(CONFIG_FOLDER, f"{template_name}_fields.json")

    try:
        if os.path.exists(file_path):
            os.remove(file_path)
        if os.path.exists(placeholders_path):
            os.remove(placeholders_path)
        if os.path.exists(fields_path):
            os.remove(fields_path)
        flash(f"模板 {filename} 及相關設定已刪除")
    except Exception as e:
        flash(f"刪除失敗: {str(e)}")
    return redirect(url_for('index'))

@app.route('/field_settings/<template_name>', methods=['GET', 'POST'])
def field_settings(template_name):
    placeholders_path = os.path.join(CONFIG_FOLDER, f"{template_name}_placeholders.json")
    fields_path = os.path.join(CONFIG_FOLDER, f"{template_name}_fields.json")

    if not os.path.exists(placeholders_path):
        flash('找不到佔位符設定，請先上傳模板')
        return redirect(url_for('index'))

    with open(placeholders_path, 'r', encoding='utf-8') as f:
        placeholders_data = json.load(f)
    placeholders = placeholders_data.get('placeholders', [])

    if request.method == 'POST':
        selected_fields = request.form.getlist('fields')
        with open(fields_path, 'w', encoding='utf-8') as f:
            json.dump({"batch_fields": selected_fields}, f, ensure_ascii=False, indent=2)
        flash('欄位設定已儲存')
        return redirect(url_for('index'))

    # 預設已選欄位
    selected_fields = []
    if os.path.exists(fields_path):
        with open(fields_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
            selected_fields = data.get('batch_fields', [])

    return render_template('field_settings.html', template_name=template_name, placeholders=placeholders, selected_fields=selected_fields)

def extract_placeholders(docx_path):
    doc = Document(docx_path)
    placeholders = set()

    def find_placeholders_in_text(text):
        # 目前假設佔位符格式為 {{欄位}} 或直接欄位名
        import re
        # 支援 {{欄位}} 格式
        matches = re.findall(r'{{(.*?)}}', text)
        placeholders.update([m.strip() for m in matches if m.strip()])
        # 也可加入直接字串匹配，視需求調整
        # matches2 = re.findall(r'\b\w+\b', text)
        # placeholders.update(matches2)

    for p in doc.paragraphs:
        find_placeholders_in_text(p.text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    find_placeholders_in_text(p.text)

    for section in doc.sections:
        for p in section.header.paragraphs:
            find_placeholders_in_text(p.text)
        for p in section.footer.paragraphs:
            find_placeholders_in_text(p.text)

    return placeholders

if __name__ == '__main__':
    app.run(debug=True)
