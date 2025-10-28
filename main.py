import json
from datetime import datetime
from docx import Document
import os
import shutil
import sys

def generate_code(date, counter):
    year = date.strftime("%y")
    month = chr(64 + date.month)
    day_num = date.day
    first_letter = chr(64 + (day_num // 26) + (1 if day_num % 26 != 0 else 0))
    second_letter = chr(64 + (day_num % 26) if day_num % 26 != 0 else 26)

    day = first_letter + second_letter
    serial_number = f'{counter:03}'
    return f'HIYES{year}{month}{day}{serial_number}'

def process_contact(contact_name):
    if not contact_name:
        return ""  # 如果 contact_name 為空，直接返回空字符串
    
    contacts_file = 'config/contacts.json'
    
    # 檢查聯繫人文件是否存在
    if not os.path.exists(contacts_file):
        print(f"錯誤: 找不到聯繫人文件 {contacts_file}")
        return contact_name
        
    # 讀取聯繫人文件
    try:
        with open(contacts_file, 'r', encoding='utf-8') as file:
            contacts = json.load(file)
    except Exception as e:
        print(f"錯誤: 無法讀取聯繫人文件: {e}")
        return contact_name
    
    # 查找匹配的聯繫人
    for contact in contacts:
        if contact['name'] == contact_name:
            return f"{contact['name']} TEL: {contact['phone']}"
    
    # 如果沒有找到匹配的聯繫人，創建新的對象
    new_contact = {"name": contact_name, "phone": ""}
    contacts.append(new_contact)
    
    # 將更新後的聯繫人列表寫回文件
    try:
        with open(contacts_file, 'w', encoding='utf-8') as file:
            json.dump(contacts, file, ensure_ascii=False, indent=4)
    except Exception as e:
        print(f"警告: 無法更新聯繫人文件: {e}")
    
    return contact_name

def process_projects(input_projects, companies_config, template_paths, output_base_path):
    date_counters = {}
    today = datetime.now()
    today_str = today.strftime("%Y-%m-%d")
    output_path = os.path.join(output_base_path, today_str)

    if not os.path.exists(output_path):
        os.makedirs(output_path)
    
    print(f"開始處理 {len(input_projects)} 個項目")
    
    if len(input_projects) == 0:
        print("錯誤: 沒有找到任何項目，請檢查 projects.json 是否為空")
        return False
    
    # 檢查模板是否存在
    for template_path in template_paths:
        if not os.path.exists(template_path):
            print(f"錯誤: 模板文件不存在: {template_path}")
            return False
    
    processed_count = 0
    for idx, project in enumerate(input_projects):
        print(f"\n處理項目 #{idx+1}: {project.get('project_name', '未命名')}")
        
        # 檢查必要字段
        required_fields = ["project_name", "company_name", "price"]
        missing_fields = [key for key in required_fields if key not in project]
        if missing_fields:
            print(f"錯誤: 項目缺少必要字段: {', '.join(missing_fields)}")
            continue
        
        print(f"尋找公司: {project['company_name']}")
        company_info = next((c for c in companies_config if c["companyName"] == project["company_name"]), None)
        if not company_info:
            print(f"錯誤: 在 companies.json 中找不到匹配的公司: {project['company_name']}")
            print(f"可用的公司: {[c.get('companyName', '') for c in companies_config]}")
            continue
        
        print(f"找到公司信息: {company_info['companyName']}")
        
        # 檢查公司配置中的必要字段
        company_required_fields = ["companyHead", "taxID", "address"]
        company_missing_fields = [field for field in company_required_fields if field not in company_info]
        if company_missing_fields:
            print(f"錯誤: 公司 '{company_info['companyName']}' 缺少必要字段: {', '.join(company_missing_fields)}")
            continue
        
        date_str = project.get("date", today_str)
        try:
            date = datetime.strptime(date_str, "%Y-%m-%d")
        except ValueError as e:
            print(f"錯誤: 日期格式不正確: {date_str}, {str(e)}")
            continue
        
        if date_str not in date_counters:
            date_counters[date_str] = 1
        else:
            date_counters[date_str] += 1
        
        code = generate_code(date, date_counters[date_str])
        
        # 使用 contact_date 如果存在，否則使用當前日期
        try:
            now = datetime.strptime(project.get("contact_date", today_str), "%Y-%m-%d")
            now_year = now.year - 1911
            now_month = now.month
            now_day = now.day
        except ValueError as e:
            print(f"錯誤: contact_date 格式不正確: {project.get('contact_date', today_str)}, {str(e)}")
            continue

        try:
            price = float(project["price"])  # 確保 price 是數字
            untaxed = round(price / 1.05)
            taxes = round(price - untaxed)
        except (TypeError, ValueError) as e:
            print(f"錯誤: price 不是有效的數字: {project['price']}, {str(e)}")
            continue

        # 處理聯繫人信息
        processed_contacts = process_contact(project.get("contacts", ""))
        
        processed_data = {
            "project_name": project["project_name"],
            "company_name": project["company_name"],
            "contacts": processed_contacts,
            "date": date_str,
            "price": f'{price:,}',
            "untaxed": f'{untaxed:,}',
            "taxes": f'{taxes:,}',
            "company_head": company_info["companyHead"],
            "taxID": company_info["taxID"],
            "company_address": company_info["address"],
            "code": code,
            "now_year": now_year,
            "now_month": now_month,
            "now_day": now_day,
            "company_info": f'{project["company_name"]} {company_info["taxID"]}'
        }

        print(f'開始生成文檔: {project["project_name"]}')
        for template_path in template_paths:
            if 'template2' in template_path:
                doc_type = '合約'
            else:
                doc_type = '報價單'
            try:
                output_file = generate_document(template_path, processed_data, output_path, doc_type)
                print(f"成功生成文檔: {output_file}")
                processed_count += 1
            except Exception as e:
                print(f"錯誤: 生成文檔時出錯: {template_path}, {str(e)}")
    
    if processed_count == 0:
        print("\n警告: 沒有成功生成任何文檔!")
        return False
    else:
        print(f"\n處理完成! 成功生成 {processed_count} 個文檔")
        return True

def move_projects_file(output_base_path):
    today_str = datetime.now().strftime("%Y-%m-%d")
    input_folder = os.path.join(os.path.dirname(__file__), 'input')
    projects_file = os.path.join(input_folder, 'projects.json')
    
    if not os.path.exists(projects_file):
        print(f"警告: 找不到原始 projects.json 文件，無法移動")
        return
    
    # 創建目標文件夾（如果不存在）
    target_folder = os.path.join(input_folder, today_str)
    if not os.path.exists(target_folder):
        os.makedirs(target_folder)
    
    # 移動文件
    target_file = os.path.join(target_folder, 'projects.json')
    try:
        shutil.move(projects_file, target_file)
        print(f'已將 projects.json 移動到 {target_file}')
    except Exception as e:
        print(f"錯誤: 移動 projects.json 文件時出錯: {str(e)}")

def generate_document(template_path, data, output_path, doc_type):
    doc = Document(template_path)
    placeholders = {
        "project_name": data["project_name"],
        "company_name": data["company_name"],
        "contacts": data["contacts"],
        "date": data["date"],
        "price": data["price"],
        "untaxed": data["untaxed"],
        "taxes": data["taxes"],
        "company_head": data["company_head"],
        "taxID": data["taxID"],
        "company_address": data["company_address"],
        "now_year": str(data["now_year"]),
        "now_month": str(data["now_month"]),
        "now_day": str(data["now_day"]),
        "code": data["code"],
        "company_info": data["company_info"]
    }
    
    for placeholder, text in placeholders.items():
        replace_text(doc, placeholder, text)
    
    output_file = os.path.join(output_path, f'{data["project_name"]}_{doc_type}.docx')
    doc.save(output_file)
    return output_file

def replace_text(doc, placeholder, text):
    replacements_made = False
    
    def replace_in_paragraph(paragraph):
        nonlocal replacements_made
        if placeholder in paragraph.text:
            inline = paragraph.runs
            replacements = []
            for i in range(len(inline)):
                if placeholder in inline[i].text:
                    replacements.append((i, inline[i].text.replace(placeholder, text)))
                    replacements_made = True
            for i, new_text in replacements:
                inline[i].text = new_text
            full_text = ''.join([run.text for run in inline])
            if placeholder in full_text:
                new_text = full_text.replace(placeholder, text)
                for i in range(len(inline)):
                    if i == 0:
                        inline[i].text = new_text
                    else:
                        inline[i].text = ''
                replacements_made = True

    for p in doc.paragraphs:
        replace_in_paragraph(p)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_in_paragraph(paragraph)

    for section in doc.sections:
        for header in section.header.paragraphs:
            replace_in_paragraph(header)
        for footer in section.footer.paragraphs:
            replace_in_paragraph(footer)
    
    if not replacements_made:
        print(f"警告: 未找到佔位符 '{placeholder}' 或未進行任何替換")

def main():
    try:
        # 檢查必要的文件
        config_dir = os.path.join(os.path.dirname(__file__), 'config')
        if not os.path.exists(config_dir):
            print(f"錯誤: 找不到 config 目錄")
            return False
        
        companies_file = os.path.join(config_dir, 'companies.json')
        if not os.path.exists(companies_file):
            print(f"錯誤: 找不到 companies.json 文件")
            return False
        
        input_dir = os.path.join(os.path.dirname(__file__), 'input')
        if not os.path.exists(input_dir):
            print(f"錯誤: 找不到 input 目錄")
            return False
        
        projects_file = os.path.join(input_dir, 'projects.json')
        if not os.path.exists(projects_file):
            print(f"錯誤: 找不到 projects.json 文件")
            return False
        
        templates_dir = os.path.join(os.path.dirname(__file__), 'templates')
        if not os.path.exists(templates_dir):
            print(f"錯誤: 找不到 templates 目錄")
            return False
        
        # 讀取配置文件
        with open(companies_file, 'r', encoding='utf-8') as file:
            companies_config = json.load(file)
        
        with open(projects_file, 'r', encoding='utf-8') as file:
            input_projects = json.load(file)

        base_dir = os.path.dirname(__file__)
        template_paths = [
            os.path.join(base_dir, 'templates', 'template.docx'),
            os.path.join(base_dir, 'templates', 'template2.docx')
        ]
        output_base_path = os.path.join(base_dir, 'output')

        if not os.path.exists(output_base_path):
            os.makedirs(output_base_path)

        # 處理專案
        success = process_projects(input_projects, companies_config, template_paths, output_base_path)
        
        # 只有在成功生成文件時才移動 projects.json
        if success:
            move_projects_file(output_base_path)
            return True
        else:
            print("因為沒有成功生成任何文件，保留 projects.json 在原位置")
            return False
    
    except Exception as e:
        print(f"發生未預期的錯誤: {str(e)}")
        return False

if __name__ == "__main__":
    success = main()
    if not success:
        print("\n程式終止: 未能成功完成文件生成")
        sys.exit(1)
    else:
        print("\n程式成功完成")