#!/usr/bin/env python3
"""
Template Migration Script
Uploads existing templates to Firebase Storage and creates Firestore documents
"""

import os
import sys
import re
from pathlib import Path
from datetime import datetime
from typing import Dict, List, Set

# Add project root to path
project_root = Path(__file__).parent.parent
sys.path.insert(0, str(project_root))

try:
    from docx import Document
    import firebase_admin
    from firebase_admin import credentials, firestore, storage
except ImportError:
    print("❌ Required packages not installed.")
    print("Please install: pip install python-docx firebase-admin")
    sys.exit(1)


class TemplateAnalyzer:
    """Analyze Word document templates for variables"""

    STANDARD_VARIABLES = {
        'project_name', 'company_name', 'contact_name', 'price', 'date',
        'created_at', 'updated_at', 'status', 'quotation_number',
        'contract_number', 'invoice_number', 'document_number'
    }

    def __init__(self, docx_path: str):
        self.docx_path = docx_path
        self.doc = Document(docx_path)

    def extract_variables(self) -> Dict[str, List[str]]:
        """Extract all {{variable}} placeholders from document"""
        variables: Set[str] = set()

        # Search in paragraphs
        for paragraph in self.doc.paragraphs:
            variables.update(self._find_variables_in_text(paragraph.text))

        # Search in tables
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        variables.update(self._find_variables_in_text(paragraph.text))

        # Search in headers/footers
        for section in self.doc.sections:
            for paragraph in section.header.paragraphs:
                variables.update(self._find_variables_in_text(paragraph.text))
            for paragraph in section.footer.paragraphs:
                variables.update(self._find_variables_in_text(paragraph.text))

        # Categorize variables
        standard = [v for v in variables if v in self.STANDARD_VARIABLES]
        custom = [v for v in variables if v not in self.STANDARD_VARIABLES]

        return {
            'standard': sorted(standard),
            'custom': sorted(custom),
            'all': sorted(list(variables))
        }

    def _find_variables_in_text(self, text: str) -> Set[str]:
        """Find all {{variable}} patterns in text"""
        # Match {{variable_name}}
        pattern = r'\{\{([a-zA-Z0-9_]+)\}\}'
        matches = re.findall(pattern, text)
        return set(matches)


def initialize_firebase():
    """Initialize Firebase Admin SDK"""

    # Check if already initialized
    if firebase_admin._apps:
        print("✅ Firebase already initialized")
        return

    # Look for service account key
    service_account_paths = [
        project_root / 'service-account-key.json',
        project_root / 'serviceAccountKey.json',
        project_root / '.firebase' / 'service-account-key.json',
    ]

    service_account_path = None
    for path in service_account_paths:
        if path.exists():
            service_account_path = path
            break

    if not service_account_path:
        print("❌ Service account key not found.")
        print("Please download from Firebase Console and save as 'service-account-key.json'")
        sys.exit(1)

    # Initialize Firebase
    cred = credentials.Certificate(str(service_account_path))
    firebase_admin.initialize_app(cred, {
        'storageBucket': 'autodocgen-prod.firebasestorage.app'
    })

    print(f"✅ Firebase initialized with service account: {service_account_path}")


def upload_template(
    template_path: Path,
    template_name: str,
    template_type: str,
    description: str = ""
) -> Dict:
    """Upload template file and create Firestore document"""

    print(f"\n📤 Processing template: {template_name}")

    # Analyze template
    print("   🔍 Analyzing template variables...")
    analyzer = TemplateAnalyzer(str(template_path))
    variables = analyzer.extract_variables()

    print(f"   ✓ Found {len(variables['all'])} variables:")
    print(f"     - Standard: {len(variables['standard'])} ({', '.join(variables['standard'][:5])}...)")
    print(f"     - Custom: {len(variables['custom'])} ({', '.join(variables['custom'][:5]) if variables['custom'] else 'none'})")

    # Upload to Storage
    print("   📤 Uploading to Firebase Storage...")
    bucket = storage.bucket()

    # Generate unique template ID
    template_id = f"{template_type}_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
    storage_path = f"templates/{template_id}/{template_path.name}"

    blob = bucket.blob(storage_path)
    blob.upload_from_filename(str(template_path))

    # Make accessible (according to storage rules, authenticated users can read)
    file_url = f"gs://{bucket.name}/{storage_path}"
    print(f"   ✓ Uploaded to: {file_url}")

    # Create Firestore document
    print("   💾 Creating Firestore document...")
    db = firestore.client()

    template_doc = {
        'name': template_name,
        'type': template_type,
        'description': description,
        'file_path': storage_path,
        'file_url': file_url,
        'file_name': template_path.name,
        'file_size': template_path.stat().st_size,
        'variables': {
            'standard': variables['standard'],
            'extra': variables['custom']
        },
        'is_active': True,
        'created_by': 'migration_script',
        'created_at': firestore.SERVER_TIMESTAMP,
        'updated_at': firestore.SERVER_TIMESTAMP,
        'shared_with': {},
        'usage_count': 0,
        'version': 1
    }

    doc_ref = db.collection('templates').document(template_id)
    doc_ref.set(template_doc)

    print(f"   ✓ Created template document: {template_id}")

    return {
        'id': template_id,
        'name': template_name,
        'variables': variables,
        'storage_path': storage_path
    }


def migrate_templates():
    """Main migration function"""

    print("=" * 60)
    print("🔥 AutoDocGen Template Migration")
    print("=" * 60)

    # Initialize Firebase
    initialize_firebase()

    # Define templates to migrate
    templates_dir = project_root / 'templates'

    templates_config = [
        {
            'path': templates_dir / 'template.docx',
            'name': 'Standard Quotation Template',
            'type': 'quote',
            'description': 'Default quotation template for HIYES projects'
        },
        {
            'path': templates_dir / 'template2.docx',
            'name': 'Alternative Template',
            'type': 'custom',
            'description': 'Secondary template for special use cases'
        }
    ]

    # Migrate each template
    migrated = []
    failed = []

    for template_config in templates_config:
        try:
            if not template_config['path'].exists():
                print(f"\n⚠️  Template not found: {template_config['path']}")
                failed.append(template_config['name'])
                continue

            result = upload_template(
                template_path=template_config['path'],
                template_name=template_config['name'],
                template_type=template_config['type'],
                description=template_config['description']
            )

            migrated.append(result)

        except Exception as e:
            print(f"\n❌ Error migrating {template_config['name']}: {e}")
            failed.append(template_config['name'])

    # Summary
    print("\n" + "=" * 60)
    print("📊 Migration Summary")
    print("=" * 60)
    print(f"✅ Successfully migrated: {len(migrated)} templates")
    for template in migrated:
        print(f"   - {template['name']} (ID: {template['id']})")
        print(f"     Variables: {len(template['variables']['all'])} total")

    if failed:
        print(f"\n❌ Failed: {len(failed)} templates")
        for name in failed:
            print(f"   - {name}")

    print("\n✨ Migration completed!")
    print("\nNext steps:")
    print("1. Verify templates in Firebase Console")
    print("2. Test template download and variable replacement")
    print("3. Update frontend to use new template IDs")


if __name__ == '__main__':
    try:
        migrate_templates()
    except KeyboardInterrupt:
        print("\n\n⚠️  Migration cancelled by user")
        sys.exit(1)
    except Exception as e:
        print(f"\n\n❌ Migration failed: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)
